from fastapi import FastAPI, HTTPException
from pydantic import BaseModel
from typing import List
from pathlib import Path
import shutil
import re
import docx
import spacy
import dateutil.parser
from sentence_transformers import SentenceTransformer
from sklearn.metrics.pairwise import cosine_similarity
import numpy as np
from fastapi.middleware.cors import CORSMiddleware
import base64
import io
from PyPDF2 import PdfReader, PdfWriter  # Changed imports


# Load NLP models
try:
    nlp = spacy.load("en_core_web_sm")
except OSError:
    from spacy.cli import download
    download("en_core_web_sm")
    nlp = spacy.load("en_core_web_sm")

sbert = SentenceTransformer("all-MiniLM-L6-v2")

# App setup
app = FastAPI()

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # Replace "*" with specific origins in production
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

UPLOAD_DIR = Path("uploads")
UPLOAD_DIR.mkdir(exist_ok=True)

# Taxonomy
TAXONOMY = {
    "Legal": ["agreement", "clause", "party", "contract", "nda"],
    "HR Policy": ["leave", "policy", "vacation", "employee", "benefits"],
    "Invoice": ["invoice", "total", "amount", "due", "bill"],
    "Meeting Note": ["minutes", "attendees", "discussion", "action items"],
    "Proposal": ["proposal", "solution", "services", "quote"],
    "Report": ["report", "summary", "annual", "metrics"],
    "JD": ["job description", "responsibilities", "skills", "qualifications"],
    "Manual": ["instructions", "troubleshooting", "usage", "guide"],
    "Research": ["abstract", "methodology", "conclusion", "study"],
    "Press Release": ["announcement", "launch", "press", "statement"],
}

# Pydantic models
class TagResponse(BaseModel):
    filename: str
    document_type: str
    taxonomy: List[str]
    date: str

class FilePayload(BaseModel):
    filename: str
    content_base64: str

# === UTILITIES ===

def extract_text_from_file(path: Path) -> str:
    if path.suffix == ".pdf":
        try:
            with open(path, 'rb') as f: # open in binary read mode for PyPDF2
                reader = PdfReader(f)
                text = ""
                for page in reader.pages:
                    text += page.extract_text() + "\n"
                return text
        except Exception as e:
            print(f"Error processing PDF: {e}")
            return ""

    elif path.suffix == ".docx":
        doc = docx.Document(path)
        return "\n".join(p.text for p in doc.paragraphs)
    elif path.suffix == ".txt":
        return path.read_text()
    return ""

def extract_date(text: str) -> str:
    patterns = [
        r"\b\d{1,2}[/-]\d{1,2}[/-]\d{2,4}\b",
        r"\b(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\s*\d{1,2},?\s*\d{4}\b",
    ]
    for pattern in patterns:
        match = re.search(pattern, text, re.IGNORECASE)
        if match:
            return match.group(0)
    return "Not Found"

def extract_noun_phrases(text: str, min_words=2, max_words=3) -> List[str]:
    doc = nlp(text)
    phrases = set()
    for chunk in doc.noun_chunks:
        phrase = chunk.text.strip()
        words = phrase.split()
        clean_words = [w for w in words if w.lower() not in nlp.Defaults.stop_words]
        if min_words <= len(clean_words) <= max_words:
            cleaned = " ".join(clean_words)
            if not re.search(r"\b\d{3,}\b|@|\.(com|net|org)", cleaned, re.I):
                phrases.add(cleaned.lower())
    return list(phrases)

def classify_type(keywords: List[str]) -> str:
    for label, vocab in TAXONOMY.items():
        if any(term.lower() in phrase.lower() for phrase in keywords for term in vocab):
            return label
    return "Unknown"

def rank_phrases_by_similarity(phrases: List[str], doc_text: str, top_n: int = 3) -> List[str]:
    if not phrases:
        return []
    phrase_embeddings = sbert.encode(phrases, convert_to_tensor=True)
    doc_embedding = sbert.encode(doc_text, convert_to_tensor=True)
    sims = cosine_similarity(
        phrase_embeddings.cpu().numpy(), doc_embedding.unsqueeze(0).cpu().numpy()
    ).flatten()
    ranked = sorted(zip(phrases, sims), key=lambda x: x[1], reverse=True)
    final_phrases = []
    seen = set()
    for phrase, _ in ranked:
        title_phrase = phrase.title().strip()
        if title_phrase not in seen:
            seen.add(title_phrase)
            final_phrases.append(title_phrase)
        if len(final_phrases) >= top_n:
            break
    return final_phrases

def refine_taxonomy_with_sbert(tags: List[str], similarity_threshold=0.8) -> List[str]:
    if not tags:
        return []
    tag_embeddings = sbert.encode(tags, convert_to_tensor=True).cpu().numpy()
    used = set()
    refined_tags = []
    for i, tag in enumerate(tags):
        if i in used:
            continue
        group = [tag]
        used.add(i)
        for j in range(i + 1, len(tags)):
            if j in used:
                continue
            sim = cosine_similarity([tag_embeddings[i]], [tag_embeddings[j]])[0][0]
            if sim >= similarity_threshold:
                group.append(tags[j])
                used.add(j)
        representative = min(group, key=lambda x: len(x))
        refined_tags.append(representative.title().strip())
    return refined_tags

# === ENDPOINT ===

@app.post("/process/combined/base64/", response_model=TagResponse)
async def process_combined(payload: FilePayload):
    try:
        file_path = UPLOAD_DIR / payload.filename

        # Validate supported file type
        if not file_path.suffix in [".pdf", ".docx", ".txt"]:
            raise HTTPException(status_code=400, detail="Unsupported file type. Use PDF, DOCX or TXT.")

        # Decode base64
        try:
            decoded_content = base64.b64decode(payload.content_base64)
        except Exception as e:
            raise HTTPException(status_code=400, detail=f"Invalid base64 string: {e}")


        # Determine file type and process accordingly without writing to disk
        if file_path.suffix == ".pdf":
            try:
                with io.BytesIO(decoded_content) as f:
                    # Attempt to fix the PDF by writing and reading it again
                    try:
                        pdf_writer = PdfWriter()
                        pdf_reader = PdfReader(f)
                        for page in pdf_reader.pages:
                            pdf_writer.add_page(page)

                        output_stream = io.BytesIO()
                        pdf_writer.write(output_stream)
                        output_stream.seek(0)  # Reset stream position to the beginning
                        pdf_reader = PdfReader(output_stream)  # Read from the fixed stream

                        text = ""
                        for page in pdf_reader.pages:
                            text += page.extract_text() + "\n"

                    except Exception as fix_e:
                        print(f"Error fixing PDF: {fix_e}")
                        f.seek(0)  # Reset stream position in case fix fails
                        pdf_reader = PdfReader(f) # Try reading as is
                        text = ""
                        for page in pdf_reader.pages:
                            text += page.extract_text() + "\n"


            except Exception as e:
                 raise HTTPException(status_code=400, detail=f"Error processing PDF: {e}")


        elif file_path.suffix == ".docx":
            try:
                with io.BytesIO(decoded_content) as f:
                    doc = docx.Document(f)
                    text = "\n".join(p.text for p in doc.paragraphs)
            except Exception as e:
                 raise HTTPException(status_code=400, detail=f"Error processing DOCX: {e}")


        elif file_path.suffix == ".txt":
            try:
                text = decoded_content.decode('utf-8')
            except Exception as e:
                raise HTTPException(status_code=400, detail=f"Error decoding TXT: {e}")

        else:
            raise HTTPException(status_code=400, detail="Unsupported file type.")


        if not text.strip():
            raise HTTPException(status_code=400, detail="Document contains no readable text.")

        noun_phrases = extract_noun_phrases(text)
        document_type = classify_type(noun_phrases)
        ranked_tags = rank_phrases_by_similarity(noun_phrases, text)
        refined_tags = refine_taxonomy_with_sbert(ranked_tags)
        date = extract_date(text)

        return TagResponse(
            filename=payload.filename,
            document_type=document_type,
            taxonomy=refined_tags,
            date=date,
        )

    except HTTPException as http_ex:
        raise http_ex  # Re-raise HTTPExceptions to preserve their status code

    except Exception as e:
        print(f"General error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

    finally:
        # No need to clean up files, as they are not written to disk
        pass